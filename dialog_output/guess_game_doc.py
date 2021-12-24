"""
generate docx for image-guessing game.
"""
import json
import random
from docx import Document
from docx.shared import Inches, Pt, Cm
import cv2
import matplotlib.pyplot as plt
from tqdm import tqdm
import argparse


def run():
    parser = argparse.ArgumentParser()
    parser.add_argument('-input_file', default=None, type=str, required=True)
    parser.add_argument('-output_file', default=None, type=str, required=True)
    parser.add_argument('-img_pool_json', default=None, type=str, required=True)
    parser.add_argument('-split', default=0, type=int)

    args = parser.parse_args()
    with open(args.input_file) as f:
        d = json.load(f)
    
    d = d['data']
    index = [i for i in range(len(d))]
    
    random.seed(32)
    random.shuffle(index)  # random permutation
    index = index[(args.split-1)*50, args.split*50]

    document = Document()
    document.styles['Normal'].font.size = Pt(10.5)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    
    img_pools = []

    for ii, i in enumerate(index):
        episode = d[i]
        img_id = episode['image_id']
        img_pool = episode['img_pool']
        random.shuffle(img_pool)
        document.add_heading("id: {0}".format(ii + 1 - s), 3)
        

        for j, img_path in enumerate(img_pool):
            img = cv2.imread(img_path)
            img = cv2.cvtColor(img,cv2.COLOR_BGR2RGB)
            plt.subplot(4, 4, j + 1)
            plt.imshow(img)
            plt.title(str(j + 1))
            plt.axis('off')
        
        plt.savefig("tmp.jpg", dpi=500, bbox_inches='tight')

        document.add_picture('tmp.jpg', width=Inches(7))
        para = document.add_paragraph()
        dialog = episode['dialog']
        for fact in dialog:
            question = 'Q: ' + fact['question'] 
            answer = 'A: ' + fact['answer']
            para.add_run(question + '               ' + answer + '\n') 
        
        document.add_page_break()
        plt.clf()

        ep_img_pool = {'img_id': img_id, 'img_pool': img_pool, 'doc_id': ii + 1 - s}
        img_pools += [ep_img_pool]

        k = img_pool.index(img_id)
        print(k+1)

    document.save(args.output_file)

    with open(args.img_pool_json, 'w') as f:
        json.dump(img_pools, f)

if __name__ == '__main__':
    run()