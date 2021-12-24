"""
generate dialog instances in docx
"""
import json
import random
from docx import Document
from docx.shared import Inches, Pt, Cm
import cv2
import matplotlib.pyplot as plt
from tqdm import tqdm
import argparse

def run():
    parser = argparse.ArgumentParser()
    parser.add_argument('-input_file', default=None, type=str, required=True)
    parser.add_argument('-output_file', default=None, type=str, required=True)
    parser.add_argument('-human_study_imgid', default=None, type=str, required=True)
    parser.add_argument('-split', default=0, type=int)

    args = parser.parse_args()
    with open(args.input_file) as f:
        d = json.load(f)

    d = d['data']
    index = [i for i in range(len(d))]
    random.seed(32)
    random.shuffle(index)  # random permutation
    index = index[(args.split-1)*50, args.split*50]

    document = Document()
    document.styles['Normal'].font.size = Pt(10.5)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    
    human_study_imgid = []

    for ii, i in tqdm(enumerate(index)):
        episode = d[i]
        img_id = episode['image_id']
        document.add_heading("id: {0}".format(ii + 1), 3)
        
        document.add_picture(img_id, width=Inches(4))
        para = document.add_paragraph()
        dialog = episode['dialog']
        para.add_run('Caption: ' + episode['caption'] + '\n\n')
        for fact in dialog:
            question = 'Q: ' + fact['question'] 
            answer = 'A: ' + fact['answer']
            para.add_run(question + '               ' + answer + '\n') 

        document.add_page_break()
        plt.clf()

        human_study_imgid += [{'img_id': img_id, 'doc_id': ii + 1}]

    document.save(args.output_file)

    with open(args.human_study_imgid, 'w') as f:
        json.dump(human_study_imgid, f)

if __name__ == '__main__':
    run()